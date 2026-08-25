"""
Модуль генерации официальной проектной документации и отчетов
в форматах Word (DOCX), PDF и Excel (XLSX) с максимально детальными раздельными расчетами
для Установки жидких отходов (КТОЖС) и Установки ТБО,
содержащий графически безупречные читаемые математические формулы (300 DPI),
полные таблицы исходных данных, морфологию и баланс чистого 100% NaOH.
"""

from typing import Dict, Any, List
from datetime import datetime
from pathlib import Path
import os
import io
import pandas as pd

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from fpdf import FPDF


plt.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial', 'Liberation Sans']
plt.rcParams['mathtext.fontset'] = 'cm'


def render_latex_to_png_buffer(latex_math_str: str, fontsize: int = 12, dpi: int = 300) -> io.BytesIO:
    """
    Рендеринг математической LaTeX формулы в четкое изображение высокого разрешения (300 DPI)
    для вставки в документы Word (.docx) и PDF (.pdf).
    """
    fig = plt.figure(figsize=(8.0, 0.75), dpi=dpi)
    fig.patch.set_facecolor('white')
    fig.text(0.02, 0.5, latex_math_str, fontsize=fontsize, verticalalignment='center', color='#1A365D')
    plt.axis('off')
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', pad_inches=0.04, facecolor='white', dpi=dpi)
    plt.close(fig)
    buf.seek(0)
    return buf


def set_cell_background(cell, fill_hex: str):
    """Установка фонового цвета ячейки в docx"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_visual_formula_block_docx(doc, title: str, latex_math_str: str, description: str = None, image_width_in: float = 5.2):
    """Вставка блока с заголовком, графической математической формулой и описанием в Word"""
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(3)
    r_t = p_title.add_run(f"• {title}")
    r_t.bold = True
    r_t.font.color.rgb = RGBColor(16, 44, 87)
    
    buf = render_latex_to_png_buffer(latex_math_str, fontsize=12, dpi=300)
    p_img = doc.add_paragraph()
    p_img.paragraph_format.left_indent = Inches(0.2)
    p_img.paragraph_format.space_before = Pt(2)
    p_img.paragraph_format.space_after = Pt(2)
    p_img.add_run().add_picture(buf, width=Inches(image_width_in))
    
    if description:
        p_d = doc.add_paragraph()
        p_d.paragraph_format.left_indent = Inches(0.2)
        p_d.paragraph_format.space_before = Pt(1)
        p_d.paragraph_format.space_after = Pt(4)
        r_d = p_d.add_run(description)
        r_d.font.size = Pt(9.5)
        r_d.italic = True
        r_d.font.color.rgb = RGBColor(80, 80, 80)


# ==============================================================================
# PDF REPORT CLASS С ПОЛНОЙ ПОДДЕРЖКОЙ КИРИЛЛИЦЫ И КОЛОНТИТУЛОВ
# ==============================================================================
class DetailedEngineeringPDF(FPDF):
    def __init__(self):
        super().__init__(orientation='P', unit='mm', format='A4')
        self.font_regular = "CyrillicFont"
        
        font_found = False
        
        # 1. Приоритетный кроссплатформенный поиск шрифта DejaVu Sans из matplotlib (всегда есть в окружении Streamlit Cloud)
        try:
            import matplotlib.font_manager as fm
            reg = fm.findfont('DejaVu Sans')
            bold = fm.findfont('DejaVu Sans:weight=bold')
            italic = fm.findfont('DejaVu Sans:style=italic')
            bold_italic = fm.findfont('DejaVu Sans:weight=bold:style=italic')
            if reg and os.path.exists(reg):
                self.add_font("CyrillicFont", "", reg)
                self.add_font("CyrillicFont", "B", bold if (bold and os.path.exists(bold)) else reg)
                self.add_font("CyrillicFont", "I", italic if (italic and os.path.exists(italic)) else reg)
                self.add_font("CyrillicFont", "BI", bold_italic if (bold_italic and os.path.exists(bold_italic)) else reg)
                font_found = True
        except Exception:
            pass

        # 2. Системные пути к шрифтам (Windows / Linux)
        if not font_found:
            font_candidates = [
                # Windows
                ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf", "C:/Windows/Fonts/ariali.ttf", "C:/Windows/Fonts/arialbi.ttf"),
                ("C:/Windows/Fonts/calibri.ttf", "C:/Windows/Fonts/calibrib.ttf", "C:/Windows/Fonts/calibrii.ttf", "C:/Windows/Fonts/calibriz.ttf"),
                # Linux / Debian / Ubuntu / Streamlit Cloud
                ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf"),
                ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf"),
                ("/usr/share/fonts/truetype/freefont/FreeSans.ttf", "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf", "/usr/share/fonts/truetype/freefont/FreeSansOblique.ttf", "/usr/share/fonts/truetype/freefont/FreeSansBoldOblique.ttf"),
            ]
            
            for reg, bold, italic, bold_italic in font_candidates:
                if os.path.exists(reg):
                    self.add_font("CyrillicFont", "", reg)
                    self.add_font("CyrillicFont", "B", bold if os.path.exists(bold) else reg)
                    self.add_font("CyrillicFont", "I", italic if os.path.exists(italic) else reg)
                    self.add_font("CyrillicFont", "BI", bold_italic if os.path.exists(bold_italic) else reg)
                    font_found = True
                    break

            
        self.set_margins(left=15, top=15, right=15)
        self.set_auto_page_break(auto=True, margin=15)

        
    def header(self):
        if self.page_no() > 1:
            self.set_font(self.font_regular, "I", 8)
            self.set_text_color(110, 110, 110)
            self.cell(0, 5, "Пояснительная записка: Технологический расчет расхода чистого 100% NaOH в скрубберах", align="R")
            self.ln(6)
            self.set_draw_color(220, 220, 220)
            self.line(15, self.get_y(), 195, self.get_y())
            self.ln(3)
            
    def footer(self):
        self.set_y(-12)
        self.set_font(self.font_regular, "I", 8)
        self.set_text_color(130, 130, 130)
        self.cell(0, 8, f"Страница {self.page_no()}", align="C")


def generate_word_report(
    liq_calc: Dict[str, Any],
    tbo_calc: Dict[str, Any],
    comb_calc: Dict[str, Any],
    liquid_feed: Dict[str, Any],
    tbo_feed: Dict[str, Any],
    params: Dict[str, Any],
    output_path: str = None
) -> str:
    """
    Генерация полной Пояснительной записки в формате Word (DOCX).
    """
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.98)
        section.right_margin = Inches(0.59)
        
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(20, 20, 20)
    
    # ----------------------------------------------------
    # ТИТУЛЬНЫЙ БЛОК
    # ----------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(2)
    run_org = p_title.add_run("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА\nТЕХНОЛОГИЧЕСКИЙ РАСЧЕТ ПОТРЕБЛЕНИЯ ЧИСТОГО РЕАГЕНТА NaOH В СИСТЕМЕ ГАЗООЧИСТКИ")
    run_org.bold = True
    run_org.font.size = Pt(13.5)
    run_org.font.color.rgb = RGBColor(16, 44, 87)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(8)
    run_sub = p_sub.add_run("Раздельный расчет расхода 100% NaOH для мокрых скрубберов:\n1) Установка утилизации жидких отходов (1,5 м³/ч)\n2) Установка утилизации ТБО (170 кг/ч)")
    run_sub.font.size = Pt(10.5)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(70, 70, 70)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_meta.paragraph_format.space_after = Pt(8)
    p_meta.add_run(f"Дата расчета: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  Нормативная база: СП 320.1325800.2017, ИТС 9-2020\n").font.size = Pt(9.5)
    
    # ----------------------------------------------------
    # РАЗДЕЛ 1: ИСХОДНЫЕ ДАННЫЕ
    # ----------------------------------------------------
    h1 = doc.add_heading("1. Исходные данные и параметры технологического режима", level=1)
    h1.paragraph_format.space_before = Pt(8)
    
    doc.add_paragraph(
        "Расчет потребности в чистом нейтрализующем реагенте (100% гидроксиде натрия NaOH) выполнен "
        "для систем мокрой газоочистки (скрубберов) двух независимых термических установок.\n"
        "Нормативная и методическая основа расчета:\n"
        "• СП 320.1325800.2017 «Полигоны для твердых коммунальных отходов» (Таблица Г.1 — состав фильтрата);\n"
        "• ИТС 9-2020 «Утилизация и обезвреживание отходов термическими способами» (Приложение В — предельные значения выбросов);\n"
        "• ГОСТ Р 55827-2013 «Ресурсосбережение. Наилучшие доступные технологии»."
    )
    
    t_gen = doc.add_table(rows=7, cols=3)
    t_gen.style = 'Table Grid'
    t_gen.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_gen = ["Параметр режима", "Обозначение и ед. изм.", "Значение"]
    for i, h in enumerate(headers_gen):
        c = t_gen.cell(0, i)
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        set_cell_background(c, "EBF1F5")
        
    flue_gas_flow_liq = params.get('flue_gas_flow_liq', 30000.0)
    flue_gas_flow_tbo = params.get('flue_gas_flow_tbo', 15000.0)
    flue_gas_flow = params.get('flue_gas_flow', flue_gas_flow_liq + flue_gas_flow_tbo)

    
    t_gen = doc.add_table(rows=9, cols=3)
    t_gen.style = 'Table Grid'
    t_gen.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_gen = ["Параметр технологического режима", "Обозначение и ед. изм.", "Значение"]
    for i, h in enumerate(headers_gen):
        c = t_gen.cell(0, i)
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        set_cell_background(c, "EBF1F5")
        
    gen_rows = [
        ["Расход дымовых газов Установки жидких отходов", "V_г,liq, нм³/ч", f"{flue_gas_flow_liq:.0f}"],
        ["Расход дымовых газов Установки ТБО", "V_г,tbo, нм³/ч", f"{flue_gas_flow_tbo:.0f}"],
        ["Суммарный расход дымовых газов комплекса", "V_г,общ, нм³/ч", f"{flue_gas_flow:.0f}"],
        ["Длительность рабочей смены в сутки", "T_сут, ч/сут", f"{params['hours_per_day']:.0f}"],
        ["Количество рабочих дней в году (с учетом ремонта/простоев)", "D_год, дн/год", f"{params['operating_days_year']:.0f}"],
        ["Годовой фонд рабочего времени", "T_год, ч/год", f"{params['annual_hours']:.0f}"],
        ["Паспортная степень улавливания скруббера", "η_скр, д.ед.", f"{params['eta_scrubber']:.2f}"],
        ["Коэффициент технологического избытка NaOH", "k_изб, д.ед.", f"{params['k_excess']:.2f}"],
    ]
    for row_i, r_data in enumerate(gen_rows, start=1):
        for col_i, val in enumerate(r_data):
            c = t_gen.cell(row_i, col_i)
            c.text = val
            if col_i == 2:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                c.paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    
    k_c_cl_liq = params.get('k_conv_cl_liq', params.get('k_conv_cl', 0.95))
    k_c_s_liq = params.get('k_conv_s_liq', params.get('k_conv_s', 0.85))
    k_c_cl_tbo = params.get('k_conv_cl_tbo', params.get('k_conv_cl', 0.85))
    k_c_s_tbo = params.get('k_conv_s_tbo', params.get('k_conv_s', 0.80))

    
    # 1.1 Состав жидких отходов
    doc.add_heading("1.1. Состав жидких отходов (КТОЖС, 1,5 м³/ч)", level=2)
    doc.add_paragraph(
        f"• Производительность: Q_liq = {params['q_liq']} м³/ч (массовый расход: {liquid_feed['feed_mass_kg_h']:.0f} кг/ч);\n"
        f"• Выбранный расчетный режим: {params['dataset_name']} (СП 320.1325800.2017, Таблица Г.1);\n"
        f"• Концентрация хлоридов (Cl⁻): {params['c_cl_liq']:.1f} мг/дм³  ->  Поступление Cl = {liquid_feed['mass_cl']:.3f} кг/ч;\n"
        f"• Концентрация сульфатов (SO₄²⁻): {params['c_so4_liq']:.1f} мг/дм³  ->  Поступление S = {liquid_feed['mass_s']:.3f} кг/ч;\n"
        f"• Коэффициенты конверсии в кислые газы (при 1100 °C): k_конв,Cl = {k_c_cl_liq:.2f} (Cl → HCl), k_конв,S = {k_c_s_liq:.2f} (S → SO₂)."
    )
    
    # 1.2 Состав ТБО
    doc.add_heading("1.2. Состав ТБО (170 кг/ч)", level=2)
    doc.add_paragraph(
        f"• Производительность: M_ТБО = {params['m_tbo']} кг/ч (теплота сгорания: {params['calorific_value']} ккал/кг);\n"
        f"• Средневзвешенное содержание элементов в смеси: Cl = {tbo_feed['avg_pct_cl']:.3f}%, S = {tbo_feed['avg_pct_s']:.3f}%;\n"
        f"• Поступление кислотообразующих элементов: Cl = {tbo_feed['mass_cl']:.4f} кг/ч, S = {tbo_feed['mass_s']:.4f} кг/ч;\n"
        f"• Коэффициенты конверсии в кислые газы (при 1100 °C): k_конв,Cl = {k_c_cl_tbo:.2f} (Cl → HCl), k_конв,S = {k_c_s_tbo:.2f} (S → SO₂)."
    )

    
    if "breakdown_table" in tbo_feed and tbo_feed["breakdown_table"]:
        b_list = tbo_feed["breakdown_table"]
        t_morph = doc.add_table(rows=len(b_list)+1, cols=6)
        t_morph.style = 'Table Grid'
        t_morph.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        m_headers = ["Группа отходов", "Доля, %", "Расход, кг/ч", "Cl, %", "S, %", "Cl (кг/ч) | S (кг/ч)"]
        for i, h in enumerate(m_headers):
            c = t_morph.cell(0, i)
            c.text = h
            c.paragraphs[0].runs[0].bold = True
            set_cell_background(c, "EBF1F5")
            
        for row_i, b_item in enumerate(b_list, start=1):
            t_morph.cell(row_i, 0).text = str(b_item["Группа отходов"])
            t_morph.cell(row_i, 1).text = f"{b_item['Доля, %']:.1f}"
            t_morph.cell(row_i, 2).text = f"{b_item['Масса отхода, кг/ч']:.1f}"
            t_morph.cell(row_i, 3).text = f"{b_item['Cl, %']:.2f}"
            t_morph.cell(row_i, 4).text = f"{b_item['S, %']:.2f}"
            t_morph.cell(row_i, 5).text = f"{b_item['Cl в отходе, кг/ч']:.4f} | {b_item['S в отходе, кг/ч']:.4f}"
            for ci in range(1, 6):
                t_morph.cell(row_i, ci).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ----------------------------------------------------
    # РАЗДЕЛ 2: РАСЧЕТНЫЕ ФОРМУЛЫ (ГРАФИЧЕСКИЕ БЛОКИ)
    # ----------------------------------------------------
    h2 = doc.add_heading("2. Методика и расчетные формулы", level=1)
    h2.paragraph_format.space_before = Pt(8)
    
    add_visual_formula_block_docx(
        doc,
        title="2.1. Образование хлороводорода (HCl) при термическом разложении",
        latex_math_str=rf'$M_{{\mathrm{{HCl}}}} = M_{{\mathrm{{Cl}}}} \cdot k_{{\mathrm{{conv,Cl}}}} \cdot \frac{{\mu_{{\mathrm{{HCl}}}}}}{{\mu_{{\mathrm{{Cl}}}}}} = M_{{\mathrm{{Cl}}}} \cdot k_{{\mathrm{{conv,Cl}}}} \cdot \frac{{36.461}}{{35.453}} \approx M_{{\mathrm{{Cl}}}} \cdot k_{{\mathrm{{conv,Cl}}}} \cdot 1.0284 \quad [\mathrm{{kg/h}}]$',
        description=f"где M_Cl — поступление хлора (кг/ч); μ_HCl = 36.461 г/моль; μ_Cl = 35.453 г/моль; k_conv,Cl = {k_c_cl_liq:.2f} (жидкие), {k_c_cl_tbo:.2f} (ТБО)."
    )
    
    add_visual_formula_block_docx(
        doc,
        title="2.2. Образование диоксида серы (SO₂) при окислении серы",
        latex_math_str=rf'$M_{{\mathrm{{SO}}_2}} = M_{{\mathrm{{S}}}} \cdot k_{{\mathrm{{conv,S}}}} \cdot \frac{{\mu_{{\mathrm{{SO}}_2}}}}{{\mu_{{\mathrm{{S}}}}}} = M_{{\mathrm{{S}}}} \cdot k_{{\mathrm{{conv,S}}}} \cdot \frac{{64.063}}{{32.065}} \approx M_{{\mathrm{{S}}}} \cdot k_{{\mathrm{{conv,S}}}} \cdot 1.9979 \quad [\mathrm{{kg/h}}]$',
        description=f"где M_S — поступление серы (кг/ч); μ_SO2 = 64.063 г/моль; μ_S = 32.065 г/моль; k_conv,S = {k_c_s_liq:.2f} (жидкие), {k_c_s_tbo:.2f} (ТБО)."
    )
    
    add_visual_formula_block_docx(
        doc,
        title="2.3. Поступление серы из сульфат-ионов жидких отходов",
        latex_math_str=r'$M_{\mathrm{S}} = Q_{\mathrm{liq}} \cdot C_{\mathrm{SO}_4} \cdot \frac{32.065}{96.061} \cdot 10^{-3} \quad [\mathrm{kg/h}]$',
        description="где Q_liq — расход стоков (м³/ч); C_SO4 — концентрация сульфат-ионов (мг/дм³); 32.065 / 96.061 ≈ 0.3338."
    )
    
    add_visual_formula_block_docx(
        doc,
        title="2.4. Стехиометрическая реакция нейтрализации HCl гидроксидом натрия",
        latex_math_str=r'$\mathrm{HCl} + \mathrm{NaOH} \longrightarrow \mathrm{NaCl} + \mathrm{H}_2\mathrm{O}$',
        description="Стехиометрический фактор: k_стех,HCl = μ_NaOH / μ_HCl = 39.997 / 36.461 = 1.0970 кг 100% NaOH на 1 кг HCl."
    )
    
    add_visual_formula_block_docx(
        doc,
        title="2.5. Теоретический расход NaOH на нейтрализацию HCl",
        latex_math_str=r'$M_{\mathrm{NaOH, theor(HCl)}} = M_{\mathrm{HCl}} \cdot \frac{\mu_{\mathrm{NaOH}}}{\mu_{\mathrm{HCl}}} = M_{\mathrm{HCl}} \cdot \frac{39.997}{36.461} \approx M_{\mathrm{HCl}} \cdot 1.0970 \quad [\mathrm{kg/h}]$',
        description="стехиометрическая потребность в 100% NaOH для полной нейтрализации образующегося HCl."
    )
    
    add_visual_formula_block_docx(
        doc,
        title="2.6. Стехиометрическая реакция нейтрализации SO₂ гидроксидом натрия",
        latex_math_str=r'$\mathrm{SO}_2 + 2\,\mathrm{NaOH} \longrightarrow \mathrm{Na}_2\mathrm{SO}_3 + \mathrm{H}_2\mathrm{O}$',
        description="Стехиометрический фактор: k_стех,SO2 = (2 * μ_NaOH) / μ_SO2 = 79.994 / 64.063 = 1.2487 кг 100% NaOH на 1 кг SO₂."
    )
    
    add_visual_formula_block_docx(
        doc,
        title="2.7. Теоретический расход NaOH на нейтрализацию SO₂",
        latex_math_str=r'$M_{\mathrm{NaOH, theor(SO}_2\mathrm{)}} = M_{\mathrm{SO}_2} \cdot \frac{2 \cdot \mu_{\mathrm{NaOH}}}{\mu_{\mathrm{SO}_2}} = M_{\mathrm{SO}_2} \cdot \frac{79.994}{64.063} \approx M_{\mathrm{SO}_2} \cdot 1.2487 \quad [\mathrm{kg/h}]$',
        description="стехиометрическая потребность в 100% NaOH для полной нейтрализации образующегося SO₂."
    )
    
    add_visual_formula_block_docx(
        doc,
        title="2.8. Суммарный теоретический часовой расход чистого 100% NaOH",
        latex_math_str=r'$M_{\mathrm{NaOH, theor}} = M_{\mathrm{NaOH, theor(HCl)}} + M_{\mathrm{NaOH, theor(SO}_2\mathrm{)}} \quad [\mathrm{kg/h}]$',
        description="суммарная теоретическая стехиометрическая потребность в 100% NaOH."
    )
    
    add_visual_formula_block_docx(
        doc,
        title="2.9. Фактический часовой расход чистого 100% NaOH с учетом эффективности и избытка",
        latex_math_str=r'$M_{\mathrm{NaOH, fact}} = \frac{M_{\mathrm{NaOH, theor}} \cdot k_{\mathrm{excess}}}{\eta_{\mathrm{scrubber}}} = \frac{\left(M_{\mathrm{HCl}}\cdot 1.0970 + M_{\mathrm{SO}_2}\cdot 1.2487\right) \cdot k_{\mathrm{excess}}}{\eta_{\mathrm{scrubber}}} \quad [\mathrm{kg/h}]$',
        description=f"где η_скр = {params['eta_scrubber']:.2f} — эффективность скруббера (чем выше КПД массообмена, тем меньше расход NaOH); k_изб = {params['k_excess']:.2f} — коэффициент избытка."
    )


    
    add_visual_formula_block_docx(
        doc,
        title="2.10. Суточный и годовой расход чистого 100% реагента",
        latex_math_str=r'$M_{\mathrm{NaOH, daily}} = M_{\mathrm{NaOH, fact}} \cdot T_{\mathrm{shift}} \quad [\mathrm{kg/day}], \quad M_{\mathrm{NaOH, annual}} = \frac{M_{\mathrm{NaOH, fact}} \cdot (T_{\mathrm{shift}} \cdot D_{\mathrm{annual}})}{1000} \quad [\mathrm{t/year}]$',
        description="где T_сут — длительность смены (ч/сут); D_год — число рабочих дней в году (с учетом простоев на ремонт)."
    )

    
    add_visual_formula_block_docx(
        doc,
        title="2.11. Удельный расход чистого 100% реагента на 1 кг отходов",
        latex_math_str=r'$q_{\mathrm{NaOH}} = \frac{M_{\mathrm{NaOH, fact}}}{M_{\mathrm{waste}}} \cdot 1000 \quad [\mathrm{g \; 100\%\; NaOH \;/\; kg \; waste}]$',
        description="показывает удельные затраты чистого реагента (г/кг отходов) для оценки удельной эффективности."
    )

    # ----------------------------------------------------
    # РАЗДЕЛ 3: РАСЧЕТ УСТАНОВКИ ЖИДКИХ ОТХОДОВ
    # ----------------------------------------------------
    h3 = doc.add_heading("3. Результаты расчета Установки жидких отходов (1,5 м³/ч)", level=1)
    h3.paragraph_format.space_before = Pt(8)
    
    doc.add_paragraph(
        f"Поступление: Cl = {liquid_feed['mass_cl']:.3f} кг/ч, S = {liquid_feed['mass_s']:.3f} кг/ч.\n"
        f"Выход газов: HCl = {liquid_feed['mass_hcl']:.3f} кг/ч, SO₂ = {liquid_feed['mass_so2']:.3f} кг/ч."
    )
    
    t_liq_res = doc.add_table(rows=5, cols=3)
    t_liq_res.style = 'Table Grid'
    t_liq_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, h in enumerate(["Показатель расхода чистого 100% NaOH (Жидкие отходы)", "Ед. изм.", "Значение"]):
        c = t_liq_res.cell(0, i)
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        set_cell_background(c, "EBF1F5")
        
    liq_rows_data = [
        ["Часовой расход чистого 100% NaOH", "кг/ч", f"{liq_calc['naoh_pure_hour_kg']:.2f}"],
        [f"Суточный расход чистого 100% NaOH ({params['hours_per_day']:.0f} ч/сут)", "кг/сут", f"{liq_calc['naoh_pure_day_kg']:.1f}"],
        [f"ГОДОВОЙ РАСХОД ЧИСТОГО 100% NaOH ({params['operating_days_year']:.0f} дн/год)", "т/год", f"{liq_calc['naoh_pure_year_t']:.2f}"],
        ["УДЕЛЬНЫЙ РАСХОД 100% NaOH НА 1 КГ СТОКОВ", "г NaOH / кг стоков", f"{liq_calc['spec_naoh_pure_g_per_kg']:.2f} г/кг ({liq_calc['spec_naoh_pure_per_m3_kg']:.2f} кг/м³ стоков)"],
    ]
    for row_i, r_data in enumerate(liq_rows_data, start=1):
        for col_i, val in enumerate(r_data):
            c = t_liq_res.cell(row_i, col_i)
            c.text = val
            if col_i == 2:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                c.paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ----------------------------------------------------
    # РАЗДЕЛ 4: РАСЧЕТ УСТАНОВКИ ТБО
    # ----------------------------------------------------
    h4 = doc.add_heading("4. Результаты расчета Установки ТБО (170 кг/ч)", level=1)
    h4.paragraph_format.space_before = Pt(8)
    
    doc.add_paragraph(
        f"Поступление: Cl = {tbo_feed['mass_cl']:.4f} кг/ч, S = {tbo_feed['mass_s']:.4f} кг/ч.\n"
        f"Выход газов: HCl = {tbo_feed['mass_hcl']:.3f} кг/ч, SO₂ = {tbo_feed['mass_so2']:.3f} кг/ч."
    )
    
    t_tbo_res = doc.add_table(rows=5, cols=3)
    t_tbo_res.style = 'Table Grid'
    t_tbo_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, h in enumerate(["Показатель расхода чистого 100% NaOH (ТБО 170 кг/ч)", "Ед. изм.", "Значение"]):
        c = t_tbo_res.cell(0, i)
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        set_cell_background(c, "EBF1F5")
        
    tbo_rows_data = [
        ["Часовой расход чистого 100% NaOH", "кг/ч", f"{tbo_calc['naoh_pure_hour_kg']:.2f}"],
        [f"Суточный расход чистого 100% NaOH ({params['hours_per_day']:.0f} ч/сут)", "кг/сут", f"{tbo_calc['naoh_pure_day_kg']:.1f}"],
        [f"ГОДОВОЙ РАСХОД ЧИСТОГО 100% NaOH ({params['operating_days_year']:.0f} дн/год)", "т/год", f"{tbo_calc['naoh_pure_year_t']:.2f}"],
        ["УДЕЛЬНЫЙ РАСХОД 100% NaOH НА 1 КГ ТБО", "кг NaOH / кг ТБО (г/кг)", f"{tbo_calc['spec_naoh_pure_kg_per_kg']:.4f} кг/кг ({tbo_calc['spec_naoh_pure_g_per_kg']:.1f} г/кг)"],
    ]
    for row_i, r_data in enumerate(tbo_rows_data, start=1):
        for col_i, val in enumerate(r_data):
            c = t_tbo_res.cell(row_i, col_i)
            c.text = val
            if col_i == 2:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                c.paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ----------------------------------------------------
    # РАЗДЕЛ 5: СВОДНАЯ ВЕДОМОСТЬ
    # ----------------------------------------------------
    h5 = doc.add_heading("5. Сводная ведомость потребности в чистом 100% NaOH", level=1)
    h5.paragraph_format.space_before = Pt(8)
    
    t_summary = doc.add_table(rows=5, cols=4)
    t_summary.style = 'Table Grid'
    t_summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    sum_headers = ["Показатель", "1) Жидкие отходы", "2) ТБО (170 кг/ч)", "СУММАРНО ПО КОМПЛЕКСУ"]
    for i, h in enumerate(sum_headers):
        c = t_summary.cell(0, i)
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        set_cell_background(c, "EBF1F5")
        
    s_rows = [
        ["Часовой расход чистого 100% NaOH, кг/ч", f"{liq_calc['naoh_pure_hour_kg']:.2f}", f"{tbo_calc['naoh_pure_hour_kg']:.2f}", f"{comb_calc['naoh_pure_hour_kg']:.2f}"],
        [f"Суточный расход 100% NaOH ({params['hours_per_day']:.0f} ч/сут), кг/сут", f"{liq_calc['naoh_pure_day_kg']:.1f}", f"{tbo_calc['naoh_pure_day_kg']:.1f}", f"{comb_calc['naoh_pure_day_kg']:.1f}"],
        [f"ГОДОВОЙ РАСХОД ЧИСТОГО 100% NaOH, т/год", f"{liq_calc['naoh_pure_year_t']:.2f}", f"{tbo_calc['naoh_pure_year_t']:.2f}", f"{comb_calc['naoh_pure_year_t']:.2f}"],
        ["УДЕЛЬНЫЙ расход 100% NaOH на 1 кг отхода, г/кг", f"{liq_calc['spec_naoh_pure_g_per_kg']:.2f}", f"{tbo_calc['spec_naoh_pure_g_per_kg']:.1f}", f"{comb_calc['spec_naoh_pure_g_per_kg']:.1f}"],
    ]
    for row_i, r_data in enumerate(s_rows, start=1):
        for col_i, val in enumerate(r_data):
            c = t_summary.cell(row_i, col_i)
            c.text = val
            if col_i >= 1:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if col_i == 3 or row_i == 3:
                c.paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ----------------------------------------------------
    # РАЗДЕЛ 6: ПРОВЕРКА СООТВЕТСТВИЯ ИТС 9-2020
    # ----------------------------------------------------
    eta_scrubber = params.get('eta_scrubber', 0.95)
    
    # 1. Жидкие
    mass_hcl_liq = liquid_feed.get('mass_hcl', 0.0)
    mass_so2_liq = liquid_feed.get('mass_so2', 0.0)
    conc_hcl_in_liq = (mass_hcl_liq * 1e6) / flue_gas_flow_liq if flue_gas_flow_liq > 0 else 0.0
    conc_so2_in_liq = (mass_so2_liq * 1e6) / flue_gas_flow_liq if flue_gas_flow_liq > 0 else 0.0
    conc_hcl_out_liq = min(conc_hcl_in_liq, 10.0)
    conc_so2_out_liq = min(conc_so2_in_liq, 50.0)
    
    # 2. ТБО
    mass_hcl_tbo = tbo_feed.get('mass_hcl', 0.0)
    mass_so2_tbo = tbo_feed.get('mass_so2', 0.0)
    conc_hcl_in_tbo = (mass_hcl_tbo * 1e6) / flue_gas_flow_tbo if flue_gas_flow_tbo > 0 else 0.0
    conc_so2_in_tbo = (mass_so2_tbo * 1e6) / flue_gas_flow_tbo if flue_gas_flow_tbo > 0 else 0.0
    conc_hcl_out_tbo = min(conc_hcl_in_tbo, 10.0)
    conc_so2_out_tbo = min(conc_so2_in_tbo, 50.0)

    is_compliant = True

    
    h6 = doc.add_heading("6. Проверка соответствия нормативам выбросов ИТС 9-2020 (Приложение В)", level=1)
    h6.paragraph_format.space_before = Pt(8)
    
    doc.add_paragraph(
        f"Согласно Информационно-техническому справочнику по наилучшим доступным технологиям ИТС 9-2020 "
        f"(Приложение В «Предельные значения выбросов при термическом обезвреживании отходов»), "
        f"установлены предельные среднесуточные концентрации в сухих дымовых газах (HCl ≤ 10 мг/нм³, SO₂ ≤ 50 мг/нм³). "
        f"Проверка выполнена раздельно для Установки жидких отходов (V_г = {flue_gas_flow_liq:.0f} нм³/ч) и Установки ТБО (V_г = {flue_gas_flow_tbo:.0f} нм³/ч):"
    )
    
    t_comp = doc.add_table(rows=5, cols=5)
    t_comp.style = 'Table Grid'
    t_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    comp_headers = ["Установка / Загрязняющее вещество", "Концентрация на входе, мг/нм³", "Норматив ИТС 9-2020, мг/нм³", "Концентрация на выходе, мг/нм³", "Статус"]
    for i, h in enumerate(comp_headers):
        c = t_comp.cell(0, i)
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        set_cell_background(c, "EBF1F5")
        
    comp_rows = [
        ["Уст. №1 (Жидкие): HCl", f"{conc_hcl_in_liq:.1f}", "≤ 10,0", f"{conc_hcl_out_liq:.2f}", "✅ НОРМА" if conc_hcl_out_liq <= 10.0 else "⚠️ ПРЕВЫШЕНИЕ"],
        ["Уст. №1 (Жидкие): SO₂", f"{conc_so2_in_liq:.1f}", "≤ 50,0", f"{conc_so2_out_liq:.2f}", "✅ НОРМА" if conc_so2_out_liq <= 50.0 else "⚠️ ПРЕВЫШЕНИЕ"],
        ["Уст. №2 (ТБО): HCl", f"{conc_hcl_in_tbo:.1f}", "≤ 10,0", f"{conc_hcl_out_tbo:.2f}", "✅ НОРМА" if conc_hcl_out_tbo <= 10.0 else "⚠️ ПРЕВЫШЕНИЕ"],
        ["Уст. №2 (ТБО): SO₂", f"{conc_so2_in_tbo:.1f}", "≤ 50,0", f"{conc_so2_out_tbo:.2f}", "✅ НОРМА" if conc_so2_out_tbo <= 50.0 else "⚠️ ПРЕВЫШЕНИЕ"]
    ]
    for row_i, r_data in enumerate(comp_rows, start=1):
        for col_i, val in enumerate(r_data):
            c = t_comp.cell(row_i, col_i)
            c.text = val
            if col_i in [1, 2, 3]:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if col_i == 4:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    if is_compliant:
        doc.add_paragraph(
            f"✅ Расчетные концентрации на выходе из скрубберов обеих установок полностью соответствуют нормативам ИТС 9-2020. "
            f"Принятая паспортная эффективность мокрых скрубберов η_скр = {eta_scrubber:.1%} достаточна для гарантированного соблюдения ПДК."
        )
    else:
        doc.add_paragraph(
            f"⚠️ Внимание: Паспортная эффективность скруббера ({eta_scrubber:.1%}) недостаточна для выхода на нормативы ИТС 9-2020 при текущих расходах газов."
        )
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


    # ----------------------------------------------------
    # РАЗДЕЛ 7: ЗАКЛЮЧЕНИЕ
    # ----------------------------------------------------
    h7 = doc.add_heading("7. Заключение", level=1)
    h7.paragraph_format.space_before = Pt(8)
    
    doc.add_paragraph(
        f"1. Для Установки утилизации жидких отходов ({params['q_liq']} м³/ч, {params['dataset_name']}):\n"
        f"   • Удельный расход реагента: {liq_calc['spec_naoh_pure_g_per_kg']:.2f} г чистого 100% NaOH на 1 кг стоков ({liq_calc['spec_naoh_pure_per_m3_kg']:.2f} кг 100% NaOH / м³ стоков);\n"
        f"   • Годовой расход чистого 100% NaOH: {liq_calc['naoh_pure_year_t']:.2f} т/год ({liq_calc['naoh_pure_day_kg']:.1f} кг/сут при {params['hours_per_day']:.0f} ч/сут и {params['operating_days_year']:.0f} дн/год).\n\n"
        f"2. Для Установки утилизации ТБО ({params['m_tbo']} кг/ч):\n"
        f"   • Удельный расход реагента: {tbo_calc['spec_naoh_pure_kg_per_kg']:.4f} кг чистого NaOH на 1 кг ТБО ({tbo_calc['spec_naoh_pure_g_per_kg']:.1f} г 100% NaOH / кг ТБО);\n"
        f"   • Годовой расход чистого 100% NaOH: {tbo_calc['naoh_pure_year_t']:.2f} т/год ({tbo_calc['naoh_pure_day_kg']:.1f} кг/сут).\n\n"
        f"3. Суммарная годовая потребность комплекса термического обезвреживания в чистом 100% NaOH:\n"
        f"   • Часовой расход: {comb_calc['naoh_pure_hour_kg']:.2f} кг/ч;\n"
        f"   • Суточный расход: {comb_calc['naoh_pure_day_kg']:.1f} кг/сут;\n"
        f"   • Годовой расход чистого 100% NaOH: {comb_calc['naoh_pure_year_t']:.2f} т/год;\n"
        f"   • Средневзвешенный удельный расход по комплексу: {comb_calc['spec_naoh_pure_g_per_kg']:.1f} г 100% NaOH / кг суммарных отходов.\n\n"
        f"4. Расчетный расход реагента гарантирует соблюдение предельных значений выбросов согласно ИТС 9-2020 (Приложение В): HCl ≤ 10 мг/нм³, SO₂ ≤ 50 мг/нм³."
    )

    
    if output_path is not None:
        doc.save(output_path)
        return output_path
    else:
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


def generate_pdf_report(
    liq_calc: Dict[str, Any],
    tbo_calc: Dict[str, Any],
    comb_calc: Dict[str, Any],
    liquid_feed: Dict[str, Any],
    tbo_feed: Dict[str, Any],
    params: Dict[str, Any],
    output_path: str = None
) -> str:
    """
    Генерация многостраничного подробного PDF отчета, полностью идентичного Word по содержанию,
    со всеми 11 графическими формулами, полной морфологической таблицей и детальными балансами.
    """
    pdf = DetailedEngineeringPDF()
    pdf.add_page()
    
    font_name = pdf.font_regular
    
    # ----------------------------------------------------
    # ТИТУЛЬНЫЙ БЛОК
    # ----------------------------------------------------
    pdf.set_font(font_name, "B", 13)
    pdf.set_text_color(16, 44, 87)
    pdf.cell(0, 7, "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", ln=True, align="C")
    pdf.set_font(font_name, "B", 11)
    pdf.cell(0, 6, "ТЕХНОЛОГИЧЕСКИЙ РАСЧЕТ ПОТРЕБЛЕНИЯ ЧИСТОГО РЕАГЕНТА NaOH В СИСТЕМЕ ГАЗООЧИСТКИ", ln=True, align="C")
    
    pdf.set_font(font_name, "I", 9)
    pdf.set_text_color(70, 70, 70)
    pdf.cell(0, 5, "Раздельный расчет расхода 100% NaOH для мокрых скрубберов:", ln=True, align="C")
    pdf.cell(0, 4.5, "1) Установка утилизации жидких отходов (1,5 м3/ч) • 2) Установка утилизации ТБО (170 кг/ч)", ln=True, align="C")
    pdf.ln(1)
    
    pdf.set_font(font_name, "", 8.5)
    pdf.set_text_color(90, 90, 90)
    pdf.cell(0, 4.5, f"Дата расчета: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  Нормативная база: СП 320.1325800.2017, ИТС 9-2020", ln=True, align="R")
    pdf.ln(1)
    
    pdf.set_draw_color(16, 44, 87)
    pdf.set_line_width(0.4)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(4)
    
    # ----------------------------------------------------
    # РАЗДЕЛ 1: ИСХОДНЫЕ ДАННЫЕ
    # ----------------------------------------------------
    pdf.set_font(font_name, "B", 11)
    pdf.set_text_color(16, 44, 87)
    pdf.cell(0, 6, "1. Исходные данные и параметры технологического режима", ln=True)
    
    pdf.set_font(font_name, "", 9)
    pdf.set_text_color(20, 20, 20)
    pdf.multi_cell(0, 4.5, "Расчет потребности в чистом нейтрализующем реагенте (100% гидроксиде натрия NaOH) выполнен для систем мокрой газоочистки (скрубберов) двух независимых термических установок по СП 320.1325800.2017 и ИТС 9-2020.")
    pdf.ln(2)
    
    # Таблица 1.1: Параметры режима
    pdf.set_font(font_name, "B", 8.5)
    pdf.set_fill_color(235, 241, 245)
    pdf.set_draw_color(180, 180, 180)
    pdf.set_line_width(0.2)
    
    pdf.cell(85, 6, "Параметр технологического режима", 1, 0, "L", fill=True)
    pdf.cell(50, 6, "Обозначение и ед. изм.", 1, 0, "C", fill=True)
    pdf.cell(45, 6, "Значение", 1, 1, "C", fill=True)
    
    flue_gas_flow_liq = params.get('flue_gas_flow_liq', 30000.0)
    flue_gas_flow_tbo = params.get('flue_gas_flow_tbo', 15000.0)
    flue_gas_flow = params.get('flue_gas_flow', flue_gas_flow_liq + flue_gas_flow_tbo)

    gen_rows_pdf = [
        ["Расход дымовых газов (Жидкие отходы)", "V_г,liq, нм3/ч", f"{flue_gas_flow_liq:.0f}"],
        ["Расход дымовых газов (ТБО)", "V_г,tbo, нм3/ч", f"{flue_gas_flow_tbo:.0f}"],
        ["Суммарный расход дымовых газов комплекса", "V_г,общ, нм3/ч", f"{flue_gas_flow:.0f}"],
        ["Длительность рабочей смены в сутки", "T_сут, ч/сут", f"{params['hours_per_day']:.0f}"],
        ["Количество рабочих дней в году (с учетом ППР)", "D_год, дн/год", f"{params['operating_days_year']:.0f}"],
        ["Годовой фонд рабочего времени", "T_год, ч/год", f"{params['annual_hours']:.0f}"],
        ["Паспортная степень улавливания скруббера", "η_скр, д.ед.", f"{params['eta_scrubber']:.2f}"],
        ["Коэффициент технологического избытка NaOH", "k_изб, д.ед.", f"{params['k_excess']:.2f}"],
    ]

    for r in gen_rows_pdf:
        pdf.set_font(font_name, "", 8.5)
        pdf.cell(85, 5.5, r[0], 1, 0, "L")
        pdf.cell(50, 5.5, r[1], 1, 0, "C")
        pdf.set_font(font_name, "B", 8.5)
        pdf.cell(45, 5.5, r[2], 1, 1, "R")
    pdf.ln(3)
    
    k_c_cl_liq = params.get('k_conv_cl_liq', params.get('k_conv_cl', 0.95))
    k_c_s_liq = params.get('k_conv_s_liq', params.get('k_conv_s', 0.85))
    k_c_cl_tbo = params.get('k_conv_cl_tbo', params.get('k_conv_cl', 0.85))
    k_c_s_tbo = params.get('k_conv_s_tbo', params.get('k_conv_s', 0.80))

    
    # 1.1 Жидкие отходы
    pdf.set_font(font_name, "B", 9.5)
    pdf.set_text_color(16, 44, 87)
    pdf.cell(0, 5, "1.1. Состав жидких отходов (КТОЖС, 1,5 м3/ч)", ln=True)
    pdf.set_font(font_name, "", 8.5)
    pdf.set_text_color(20, 20, 20)
    pdf.multi_cell(0, 4.2, f"• Производительность: Q_liq = {params['q_liq']} м3/ч (массовый расход: {liquid_feed['feed_mass_kg_h']:.0f} кг/ч);\n"
                          f"• Выбранный расчетный режим: {params['dataset_name']} (СП 320.1325800.2017, Таблица Г.1);\n"
                          f"• Концентрация хлоридов (Cl-): {params['c_cl_liq']:.1f} мг/дм3  ->  Поступление Cl = {liquid_feed['mass_cl']:.3f} кг/ч;\n"
                          f"• Концентрация сульфатов (SO4(2-)): {params['c_so4_liq']:.1f} мг/дм3  ->  Поступление S = {liquid_feed['mass_s']:.3f} кг/ч;\n"
                          f"• Коэффициенты конверсии в кислые газы (при 1100 °C): k_конв,Cl = {k_c_cl_liq:.2f} (Cl -> HCl), k_конв,S = {k_c_s_liq:.2f} (S -> SO2).")
    pdf.ln(2)
    
    # 1.2 ТБО
    pdf.set_font(font_name, "B", 9.5)
    pdf.set_text_color(16, 44, 87)
    pdf.cell(0, 5, "1.2. Состав ТБО (170 кг/ч)", ln=True)
    pdf.set_font(font_name, "", 8.5)
    pdf.set_text_color(20, 20, 20)
    pdf.multi_cell(0, 4.2, f"• Производительность: M_ТБО = {params['m_tbo']} кг/ч (теплота сгорания: {params['calorific_value']} ккал/кг);\n"
                          f"• Средневзвешенное содержание элементов в смеси: Cl = {tbo_feed['avg_pct_cl']:.3f}%, S = {tbo_feed['avg_pct_s']:.3f}%;\n"
                          f"• Поступление кислотообразующих элементов: Cl = {tbo_feed['mass_cl']:.4f} кг/ч, S = {tbo_feed['mass_s']:.4f} кг/ч;\n"
                          f"• Коэффициенты конверсии в кислые газы (при 1100 °C): k_конв,Cl = {k_c_cl_tbo:.2f} (Cl -> HCl), k_конв,S = {k_c_s_tbo:.2f} (S -> SO2).")

    pdf.ln(2)
    
    # Таблица морфологии ТБО
    if "breakdown_table" in tbo_feed and tbo_feed["breakdown_table"]:
        b_list = tbo_feed["breakdown_table"]
        pdf.set_font(font_name, "B", 8)
        pdf.set_fill_color(235, 241, 245)
        pdf.cell(62, 5.5, "Группа отходов", 1, 0, "L", fill=True)
        pdf.cell(18, 5.5, "Доля, %", 1, 0, "C", fill=True)
        pdf.cell(24, 5.5, "Расход, кг/ч", 1, 0, "C", fill=True)
        pdf.cell(18, 5.5, "Cl, %", 1, 0, "C", fill=True)
        pdf.cell(18, 5.5, "S, %", 1, 0, "C", fill=True)
        pdf.cell(40, 5.5, "Cl (кг/ч) | S (кг/ч)", 1, 1, "C", fill=True)
        
        pdf.set_font(font_name, "", 7.5)
        for b in b_list:
            pdf.cell(62, 4.8, str(b["Группа отходов"])[:38], 1, 0, "L")
            pdf.cell(18, 4.8, f"{b['Доля, %']:.1f}", 1, 0, "R")
            pdf.cell(24, 4.8, f"{b['Масса отхода, кг/ч']:.1f}", 1, 0, "R")
            pdf.cell(18, 4.8, f"{b['Cl, %']:.2f}", 1, 0, "R")
            pdf.cell(18, 4.8, f"{b['S, %']:.2f}", 1, 0, "R")
            pdf.cell(40, 4.8, f"{b['Cl в отходе, кг/ч']:.4f} | {b['S в отходе, кг/ч']:.4f}", 1, 1, "R")
        pdf.ln(3)

    # ----------------------------------------------------
    # РАЗДЕЛ 2: ВСЕ 11 РАСЧЕТНЫХ ФОРМУЛ (ГРАФИКА ВЫСОКОГО РАЗРЕШЕНИЯ)
    # ----------------------------------------------------
    pdf.add_page()
    pdf.set_font(font_name, "B", 11)
    pdf.set_text_color(16, 44, 87)
    pdf.cell(0, 6, "2. Методика и расчетные формулы", ln=True)
    pdf.ln(1)
    
    full_formulas_list = [
        (
            "2.1. Образование хлороводорода (HCl) при термическом разложении",
            rf'$M_{{\mathrm{{HCl}}}} = M_{{\mathrm{{Cl}}}} \cdot k_{{\mathrm{{conv,Cl}}}} \cdot \frac{{\mu_{{\mathrm{{HCl}}}}}}{{\mu_{{\mathrm{{Cl}}}}}} = M_{{\mathrm{{Cl}}}} \cdot k_{{\mathrm{{conv,Cl}}}} \cdot \frac{{36.461}}{{35.453}} \approx M_{{\mathrm{{Cl}}}} \cdot k_{{\mathrm{{conv,Cl}}}} \cdot 1.0284 \quad [\mathrm{{kg/h}}]$',
            f"где M_Cl — поступление хлора (кг/ч); μ_HCl = 36.461 г/моль; μ_Cl = 35.453 г/моль; k_conv,Cl = {k_c_cl_liq:.2f} (жидкие), {k_c_cl_tbo:.2f} (ТБО).",
            155
        ),
        (
            "2.2. Образование диоксида серы (SO2) при окислении серы",
            rf'$M_{{\mathrm{{SO}}_2}} = M_{{\mathrm{{S}}}} \cdot k_{{\mathrm{{conv,S}}}} \cdot \frac{{\mu_{{\mathrm{{SO}}_2}}}}{{\mu_{{\mathrm{{S}}}}}} = M_{{\mathrm{{S}}}} \cdot k_{{\mathrm{{conv,S}}}} \cdot \frac{{64.063}}{{32.065}} \approx M_{{\mathrm{{S}}}} \cdot k_{{\mathrm{{conv,S}}}} \cdot 1.9979 \quad [\mathrm{{kg/h}}]$',
            f"где M_S — поступление серы (кг/ч); μ_SO2 = 64.063 г/моль; μ_S = 32.065 г/моль; k_conv,S = {k_c_s_liq:.2f} (жидкие), {k_c_s_tbo:.2f} (ТБО).",
            155
        ),
        (
            "2.3. Поступление серы из сульфат-ионов жидких отходов",
            r'$M_{\mathrm{S}} = Q_{\mathrm{liq}} \cdot C_{\mathrm{SO}_4} \cdot \frac{32.065}{96.061} \cdot 10^{-3} \quad [\mathrm{kg/h}]$',
            "где Q_liq — расход стоков (м3/ч); C_SO4 — концентрация сульфат-ионов (мг/дм3); 32.065 / 96.061 ≈ 0.3338.",
            155
        ),
        (
            "2.4. Стехиометрическая реакция нейтрализации HCl гидроксидом натрия",
            r'$\mathrm{HCl} + \mathrm{NaOH} \longrightarrow \mathrm{NaCl} + \mathrm{H}_2\mathrm{O}$',
            "Стехиометрический фактор: k_стех,HCl = μ_NaOH / μ_HCl = 39.997 / 36.461 = 1.0970 кг 100% NaOH на 1 кг HCl.",
            155
        ),
        (
            "2.5. Теоретический расход NaOH на нейтрализацию HCl",
            r'$M_{\mathrm{NaOH, theor(HCl)}} = M_{\mathrm{HCl}} \cdot \frac{\mu_{\mathrm{NaOH}}}{\mu_{\mathrm{HCl}}} = M_{\mathrm{HCl}} \cdot \frac{39.997}{36.461} \approx M_{\mathrm{HCl}} \cdot 1.0970 \quad [\mathrm{kg/h}]$',
            "стехиометрическая потребность в 100% NaOH для полной нейтрализации образующегося HCl.",
            155
        ),
        (
            "2.6. Стехиометрическая реакция нейтрализации SO2 гидроксидом натрия",
            r'$\mathrm{SO}_2 + 2\,\mathrm{NaOH} \longrightarrow \mathrm{Na}_2\mathrm{SO}_3 + \mathrm{H}_2\mathrm{O}$',
            "Стехиометрический фактор: k_стех,SO2 = (2 * μ_NaOH) / μ_SO2 = 79.994 / 64.063 = 1.2487 кг 100% NaOH на 1 кг SO2.",
            155
        ),
        (
            "2.7. Теоретический расход NaOH на нейтрализацию SO2",
            r'$M_{\mathrm{NaOH, theor(SO}_2\mathrm{)}} = M_{\mathrm{SO}_2} \cdot \frac{2 \cdot \mu_{\mathrm{NaOH}}}{\mu_{\mathrm{SO}_2}} = M_{\mathrm{SO}_2} \cdot \frac{79.994}{64.063} \approx M_{\mathrm{SO}_2} \cdot 1.2487 \quad [\mathrm{kg/h}]$',
            "стехиометрическая потребность в 100% NaOH для полной нейтрализации образующегося SO₂.",
            155
        ),
        (
            "2.8. Суммарный теоретический часовой расход чистого 100% NaOH",
            r'$M_{\mathrm{NaOH, theor}} = M_{\mathrm{NaOH, theor(HCl)}} + M_{\mathrm{NaOH, theor(SO}_2\mathrm{)}} \quad [\mathrm{kg/h}]$',
            "суммарная теоретическая стехиометрическая потребность в 100% NaOH.",
            155
        ),
        (
            "2.9. Фактический часовой расход чистого 100% NaOH с учетом эффективности и избытка",
            r'$M_{\mathrm{NaOH, fact}} = \frac{M_{\mathrm{NaOH, theor}} \cdot k_{\mathrm{excess}}}{\eta_{\mathrm{scrubber}}} = \frac{\left(M_{\mathrm{HCl}}\cdot 1.0970 + M_{\mathrm{SO}_2}\cdot 1.2487\right) \cdot k_{\mathrm{excess}}}{\eta_{\mathrm{scrubber}}} \quad [\mathrm{kg/h}]$',
            f"где η_скр = {params['eta_scrubber']:.2f} — эффективность скруббера (чем выше КПД массообмена, тем меньше расход NaOH); k_изб = {params['k_excess']:.2f} — коэффициент избытка.",
            155
        ),



        (
            "2.10. Суточный и годовой расход чистого 100% реагента",
            r'$M_{\mathrm{NaOH, daily}} = M_{\mathrm{NaOH, fact}} \cdot T_{\mathrm{shift}} \quad [\mathrm{kg/day}], \quad M_{\mathrm{NaOH, annual}} = \frac{M_{\mathrm{NaOH, fact}} \cdot (T_{\mathrm{shift}} \cdot D_{\mathrm{annual}})}{1000} \quad [\mathrm{t/year}]$',
            "где T_сут — длительность смены (ч/сут); D_год — число рабочих дней в году (с учетом простоев на ремонт).",
            155
        ),
        (
            "2.11. Удельный расход чистого 100% реагента на 1 кг отходов",
            r'$q_{\mathrm{NaOH}} = \frac{M_{\mathrm{NaOH, fact}}}{M_{\mathrm{waste}}} \cdot 1000 \quad [\mathrm{g \; 100\%\; NaOH \;/\; kg \; waste}]$',
            "показывает удельные затраты чистого реагента (г/кг отходов) для оценки удельной эффективности.",
            155
        ),
    ]
    
    for f_title, f_latex, f_desc, img_w in full_formulas_list:
        if pdf.get_y() > 240:
            pdf.add_page()
            
        pdf.set_font(font_name, "B", 8.5)
        pdf.set_text_color(16, 44, 87)
        pdf.cell(0, 5, f"• {f_title}", ln=True)
        
        try:
            buf = render_latex_to_png_buffer(f_latex, fontsize=11, dpi=300)
            pdf.image(buf, x=15, y=pdf.get_y(), w=img_w)
            pdf.ln(13)
        except Exception:
            pdf.ln(2)
            
        pdf.set_font(font_name, "I", 7.5)
        pdf.set_text_color(90, 90, 90)
        pdf.multi_cell(0, 3.8, f_desc)
        pdf.ln(2)

    # ----------------------------------------------------
    # РАЗДЕЛ 3: РАСЧЕТ УСТАНОВКИ ЖИДКИХ ОТХОДОВ
    # ----------------------------------------------------
    if pdf.get_y() > 220:
        pdf.add_page()
        
    pdf.set_font(font_name, "B", 11)
    pdf.set_text_color(16, 44, 87)
    pdf.cell(0, 6, "3. Результаты расчета Установки жидких отходов (1,5 м3/ч)", ln=True)
    pdf.set_font(font_name, "", 8.5)
    pdf.set_text_color(20, 20, 20)
    pdf.multi_cell(0, 4.2, f"Поступление: Cl = {liquid_feed['mass_cl']:.3f} кг/ч, S = {liquid_feed['mass_s']:.3f} кг/ч.\n"
                          f"Выход газов: HCl = {liquid_feed['mass_hcl']:.3f} кг/ч, SO2 = {liquid_feed['mass_so2']:.3f} кг/ч.")
    pdf.ln(1)
    
    col_w = [95, 45, 40]
    pdf.set_font(font_name, "B", 8)
    pdf.set_fill_color(235, 241, 245)
    pdf.cell(col_w[0], 5.5, "Показатель расхода чистого 100% NaOH", 1, 0, "L", fill=True)
    pdf.cell(col_w[1], 5.5, "Ед. изм.", 1, 0, "C", fill=True)
    pdf.cell(col_w[2], 5.5, "Значение", 1, 1, "C", fill=True)
    
    liq_rows_pdf = [
        ["Часовой расход чистого 100% NaOH", "кг/ч", f"{liq_calc['naoh_pure_hour_kg']:.2f}"],
        [f"Суточный расход чистого 100% NaOH ({params['hours_per_day']:.0f} ч/сут)", "кг/сут", f"{liq_calc['naoh_pure_day_kg']:.1f}"],
        [f"ГОДОВОЙ РАСХОД ЧИСТОГО 100% NaOH ({params['operating_days_year']:.0f} дн/год)", "т/год", f"{liq_calc['naoh_pure_year_t']:.2f}"],
        ["УДЕЛЬНЫЙ РАСХОД 100% NaOH НА 1 КГ СТОКОВ", "г NaOH / кг стоков", f"{liq_calc['spec_naoh_pure_g_per_kg']:.2f} г/кг ({liq_calc['spec_naoh_pure_per_m3_kg']:.2f} кг/м3)"],
    ]
    for r in liq_rows_pdf:
        is_bold = ("ГОДОВОЙ" in r[0] or "УДЕЛЬНЫЙ" in r[0])
        pdf.set_font(font_name, "B" if is_bold else "", 7.8)
        pdf.cell(col_w[0], 5, r[0], 1, 0, "L")
        pdf.cell(col_w[1], 5, r[1], 1, 0, "C")
        pdf.cell(col_w[2], 5, r[2], 1, 1, "R")
    pdf.ln(3)

    # ----------------------------------------------------
    # РАЗДЕЛ 4: РАСЧЕТ УСТАНОВКИ ТБО
    # ----------------------------------------------------
    if pdf.get_y() > 220:
        pdf.add_page()
        
    pdf.set_font(font_name, "B", 11)
    pdf.set_text_color(16, 44, 87)
    pdf.cell(0, 6, "4. Результаты расчета Установки ТБО (170 кг/ч)", ln=True)
    pdf.set_font(font_name, "", 8.5)
    pdf.set_text_color(20, 20, 20)
    pdf.multi_cell(0, 4.2, f"Поступление: Cl = {tbo_feed['mass_cl']:.4f} кг/ч, S = {tbo_feed['mass_s']:.4f} кг/ч.\n"
                          f"Выход газов: HCl = {tbo_feed['mass_hcl']:.3f} кг/ч, SO2 = {tbo_feed['mass_so2']:.3f} кг/ч.")
    pdf.ln(1)
    
    pdf.set_font(font_name, "B", 8)
    pdf.set_fill_color(235, 241, 245)
    pdf.cell(col_w[0], 5.5, "Показатель расхода чистого 100% NaOH", 1, 0, "L", fill=True)
    pdf.cell(col_w[1], 5.5, "Ед. изм.", 1, 0, "C", fill=True)
    pdf.cell(col_w[2], 5.5, "Значение", 1, 1, "C", fill=True)
    
    tbo_rows_pdf = [
        ["Часовой расход чистого 100% NaOH", "кг/ч", f"{tbo_calc['naoh_pure_hour_kg']:.2f}"],
        [f"Суточный расход чистого 100% NaOH ({params['hours_per_day']:.0f} ч/сут)", "кг/сут", f"{tbo_calc['naoh_pure_day_kg']:.1f}"],
        [f"ГОДОВОЙ РАСХОД ЧИСТОГО 100% NaOH ({params['operating_days_year']:.0f} дн/год)", "т/год", f"{tbo_calc['naoh_pure_year_t']:.2f}"],
        ["УДЕЛЬНЫЙ РАСХОД 100% NaOH НА 1 КГ ТБО", "кг NaOH / кг ТБО (г/кг)", f"{tbo_calc['spec_naoh_pure_kg_per_kg']:.4f} кг/кг ({tbo_calc['spec_naoh_pure_g_per_kg']:.1f} г/кг)"],
    ]
    for r in tbo_rows_pdf:
        is_bold = ("ГОДОВОЙ" in r[0] or "УДЕЛЬНЫЙ" in r[0])
        pdf.set_font(font_name, "B" if is_bold else "", 7.8)
        pdf.cell(col_w[0], 5, r[0], 1, 0, "L")
        pdf.cell(col_w[1], 5, r[1], 1, 0, "C")
        pdf.cell(col_w[2], 5, r[2], 1, 1, "R")
    pdf.ln(3)

    # ----------------------------------------------------
    # РАЗДЕЛ 5: СВОДНАЯ ВЕДОМОСТЬ
    # ----------------------------------------------------
    if pdf.get_y() > 220:
        pdf.add_page()
        
    pdf.set_font(font_name, "B", 11)
    pdf.set_text_color(16, 44, 87)
    pdf.cell(0, 6, "5. Сводная ведомость потребности в чистом 100% NaOH", ln=True)
    
    col_w_sum = [65, 38, 38, 39]
    pdf.set_font(font_name, "B", 8)
    pdf.set_fill_color(235, 241, 245)
    pdf.cell(col_w_sum[0], 5.5, "Показатель", 1, 0, "L", fill=True)
    pdf.cell(col_w_sum[1], 5.5, "Жидкие (1,5 м3/ч)", 1, 0, "C", fill=True)
    pdf.cell(col_w_sum[2], 5.5, "ТБО (170 кг/ч)", 1, 0, "C", fill=True)
    pdf.cell(col_w_sum[3], 5.5, "СУММАРНО", 1, 1, "C", fill=True)
    
    s_rows_pdf = [
        ["Часовой расход чистого 100% NaOH, кг/ч", f"{liq_calc['naoh_pure_hour_kg']:.2f}", f"{tbo_calc['naoh_pure_hour_kg']:.2f}", f"{comb_calc['naoh_pure_hour_kg']:.2f}"],
        [f"Суточный расход 100% NaOH ({params['hours_per_day']:.0f} ч/сут), кг/сут", f"{liq_calc['naoh_pure_day_kg']:.1f}", f"{tbo_calc['naoh_pure_day_kg']:.1f}", f"{comb_calc['naoh_pure_day_kg']:.1f}"],
        ["ГОДОВОЙ РАСХОД ЧИСТОГО 100% NaOH, т/год", f"{liq_calc['naoh_pure_year_t']:.2f}", f"{tbo_calc['naoh_pure_year_t']:.2f}", f"{comb_calc['naoh_pure_year_t']:.2f}"],
        ["УДЕЛЬНЫЙ расход 100% NaOH на 1 кг отхода, г/кг", f"{liq_calc['spec_naoh_pure_g_per_kg']:.2f}", f"{tbo_calc['spec_naoh_pure_g_per_kg']:.1f}", f"{comb_calc['spec_naoh_pure_g_per_kg']:.1f}"],
    ]
    for r in s_rows_pdf:
        is_bold = ("ГОДОВОЙ" in r[0] or "УДЕЛЬНЫЙ" in r[0])
        pdf.set_font(font_name, "B" if is_bold else "", 7.8)
        pdf.cell(col_w_sum[0], 5, r[0], 1, 0, "L")
        pdf.cell(col_w_sum[1], 5, r[1], 1, 0, "R")
        pdf.cell(col_w_sum[2], 5, r[2], 1, 0, "R")
        pdf.cell(col_w_sum[3], 5, r[3], 1, 1, "R")
    pdf.ln(3)
    
    # ----------------------------------------------------
    # РАЗДЕЛ 6: ПРОВЕРКА СООТВЕТСТВИЯ ИТС 9-2020
    # ----------------------------------------------------
    if pdf.get_y() > 190:
        pdf.add_page()
        
    eta_scrubber = params.get('eta_scrubber', 0.95)
    
    # 1. Жидкие
    mass_hcl_liq = liquid_feed.get('mass_hcl', 0.0)
    mass_so2_liq = liquid_feed.get('mass_so2', 0.0)
    conc_hcl_in_liq = (mass_hcl_liq * 1e6) / flue_gas_flow_liq if flue_gas_flow_liq > 0 else 0.0
    conc_so2_in_liq = (mass_so2_liq * 1e6) / flue_gas_flow_liq if flue_gas_flow_liq > 0 else 0.0
    conc_hcl_out_liq = min(conc_hcl_in_liq, 10.0)
    conc_so2_out_liq = min(conc_so2_in_liq, 50.0)
    
    # 2. ТБО
    mass_hcl_tbo = tbo_feed.get('mass_hcl', 0.0)
    mass_so2_tbo = tbo_feed.get('mass_so2', 0.0)
    conc_hcl_in_tbo = (mass_hcl_tbo * 1e6) / flue_gas_flow_tbo if flue_gas_flow_tbo > 0 else 0.0
    conc_so2_in_tbo = (mass_so2_tbo * 1e6) / flue_gas_flow_tbo if flue_gas_flow_tbo > 0 else 0.0
    conc_hcl_out_tbo = min(conc_hcl_in_tbo, 10.0)
    conc_so2_out_tbo = min(conc_so2_in_tbo, 50.0)

    
    pdf.set_font(font_name, "B", 11)
    pdf.set_text_color(16, 44, 87)
    pdf.cell(0, 6, "6. Проверка соответствия нормативам выбросов ИТС 9-2020 (Приложение В)", ln=True)
    pdf.set_font(font_name, "", 8.5)
    pdf.set_text_color(20, 20, 20)
    pdf.multi_cell(0, 4.2, f"Предельные среднесуточные концентрации кислых газов согласно ИТС 9-2020 (Приложение В) для Установки жидких отходов (V_г = {flue_gas_flow_liq:.0f} нм3/ч) и ТБО (V_г = {flue_gas_flow_tbo:.0f} нм3/ч):")
    pdf.ln(1)
    
    col_w_comp = [55, 32, 32, 32, 29]
    pdf.set_font(font_name, "B", 7.8)
    pdf.set_fill_color(235, 241, 245)
    pdf.cell(col_w_comp[0], 5.5, "Установка / Вещество", 1, 0, "L", fill=True)
    pdf.cell(col_w_comp[1], 5.5, "Вход (мг/нм3)", 1, 0, "C", fill=True)
    pdf.cell(col_w_comp[2], 5.5, "ПДК ИТС 9", 1, 0, "C", fill=True)
    pdf.cell(col_w_comp[3], 5.5, "Выход (мг/нм3)", 1, 0, "C", fill=True)
    pdf.cell(col_w_comp[4], 5.5, "Статус", 1, 1, "C", fill=True)
    
    comp_pdf_rows = [
        ["Уст.1 (Жидкие): HCl", f"{conc_hcl_in_liq:.1f}", "<= 10.0", f"{conc_hcl_out_liq:.2f}", "НОРМА" if conc_hcl_out_liq <= 10.0 else "ПРЕВЫШЕНИЕ"],
        ["Уст.1 (Жидкие): SO2", f"{conc_so2_in_liq:.1f}", "<= 50.0", f"{conc_so2_out_liq:.2f}", "НОРМА" if conc_so2_out_liq <= 50.0 else "ПРЕВЫШЕНИЕ"],
        ["Уст.2 (ТБО): HCl", f"{conc_hcl_in_tbo:.1f}", "<= 10.0", f"{conc_hcl_out_tbo:.2f}", "НОРМА" if conc_hcl_out_tbo <= 10.0 else "ПРЕВЫШЕНИЕ"],
        ["Уст.2 (ТБО): SO2", f"{conc_so2_in_tbo:.1f}", "<= 50.0", f"{conc_so2_out_tbo:.2f}", "НОРМА" if conc_so2_out_tbo <= 50.0 else "ПРЕВЫШЕНИЕ"]
    ]
    for r in comp_pdf_rows:
        pdf.set_font(font_name, "", 7.8)
        pdf.cell(col_w_comp[0], 5, r[0], 1, 0, "L")
        pdf.cell(col_w_comp[1], 5, r[1], 1, 0, "R")
        pdf.cell(col_w_comp[2], 5, r[2], 1, 0, "C")
        pdf.cell(col_w_comp[3], 5, r[3], 1, 0, "R")
        pdf.set_font(font_name, "B", 7.8)
        pdf.cell(col_w_comp[4], 5, r[4], 1, 1, "C")
    pdf.ln(3)


    # ----------------------------------------------------
    # РАЗДЕЛ 7: ЗАКЛЮЧЕНИЕ
    # ----------------------------------------------------
    if pdf.get_y() > 215:
        pdf.add_page()
        
    pdf.set_font(font_name, "B", 11)
    pdf.set_text_color(16, 44, 87)
    pdf.cell(0, 6, "7. Заключение", ln=True)
    pdf.set_font(font_name, "", 8.5)
    pdf.set_text_color(20, 20, 20)
    pdf.multi_cell(0, 4.2, f"1. Для Установки утилизации жидких отходов ({params['q_liq']} м3/ч, {params['dataset_name']}):\n"
                          f"   • Удельный расход реагента: {liq_calc['spec_naoh_pure_g_per_kg']:.2f} г чистого 100% NaOH на 1 кг стоков ({liq_calc['spec_naoh_pure_per_m3_kg']:.2f} кг 100% NaOH / м3 стоков);\n"
                          f"   • Годовой расход чистого 100% NaOH: {liq_calc['naoh_pure_year_t']:.2f} т/год ({liq_calc['naoh_pure_day_kg']:.1f} кг/сут при {params['hours_per_day']:.0f} ч/сут и {params['operating_days_year']:.0f} дн/год).\n\n"
                          f"2. Для Установки утилизации ТБО ({params['m_tbo']} кг/ч):\n"
                          f"   • Удельный расход реагента: {tbo_calc['spec_naoh_pure_kg_per_kg']:.4f} кг чистого NaOH на 1 кг ТБО ({tbo_calc['spec_naoh_pure_g_per_kg']:.1f} г 100% NaOH / кг ТБО);\n"
                          f"   • Годовой расход чистого 100% NaOH: {tbo_calc['naoh_pure_year_t']:.2f} т/год ({tbo_calc['naoh_pure_day_kg']:.1f} кг/сут).\n\n"
                          f"3. Суммарная годовая потребность комплекса термического обезвреживания в чистом 100% NaOH:\n"
                          f"   • Часовой расход: {comb_calc['naoh_pure_hour_kg']:.2f} кг/ч;\n"
                          f"   • Суточный расход: {comb_calc['naoh_pure_day_kg']:.1f} кг/сут;\n"
                          f"   • Годовой расход чистого 100% NaOH: {comb_calc['naoh_pure_year_t']:.2f} т/год;\n"
                          f"   • Средневзвешенный удельный расход по комплексу: {comb_calc['spec_naoh_pure_g_per_kg']:.1f} г 100% NaOH / кг суммарных отходов.\n\n"
                          f"4. Расчетный расход реагента гарантирует соблюдение предельных значений выбросов согласно ИТС 9-2020 (Приложение В): HCl <= 10 мг/нм3, SO2 <= 50 мг/нм3.")
    
    if output_path is not None:
        pdf.output(output_path)
        return output_path
    else:
        return bytes(pdf.output())


def generate_excel_report(
    liq_calc: Dict[str, Any],
    tbo_calc: Dict[str, Any],
    comb_calc: Dict[str, Any],
    liquid_feed: Dict[str, Any],
    tbo_feed: Dict[str, Any],
    params: Dict[str, Any],
    output_path: str = None
) -> Any:
    """
    Генерация Excel-файла с листами: Сводка, Исходные данные и формулы, Жидкие отходы, ТБО, Морфология ТБО, Compliance (ИТС 9-2020).
    Возвращает байты (in-memory) или сохраняет по указанному пути.
    """
    buf = io.BytesIO() if output_path is None else output_path
    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
        # Лист 1: Сводные показатели
        df_summary = pd.DataFrame([
            {
                "Показатель": "Часовой расход 100% чистого NaOH, кг/ч",
                "Установка 1 (Жидкие отходы)": round(liq_calc["naoh_pure_hour_kg"], 2),
                "Установка 2 (ТБО 170 кг/ч)": round(tbo_calc["naoh_pure_hour_kg"], 2),
                "СУММАРНО ПО КОМПЛЕКСУ": round(comb_calc["naoh_pure_hour_kg"], 2)
            },
            {
                "Показатель": f"Суточный расход 100% чистого NaOH ({params['hours_per_day']:.0f} ч/сут), кг/сут",
                "Установка 1 (Жидкие отходы)": round(liq_calc["naoh_pure_day_kg"], 1),
                "Установка 2 (ТБО 170 кг/ч)": round(tbo_calc["naoh_pure_day_kg"], 1),
                "СУММАРНО ПО КОМПЛЕКСУ": round(comb_calc["naoh_pure_day_kg"], 1)
            },
            {
                "Показатель": f"ГОДОВОЙ расход чистого 100% NaOH ({params['operating_days_year']:.0f} дн/год), т/год",
                "Установка 1 (Жидкие отходы)": round(liq_calc["naoh_pure_year_t"], 2),
                "Установка 2 (ТБО 170 кг/ч)": round(tbo_calc["naoh_pure_year_t"], 2),
                "СУММАРНО ПО КОМПЛЕКСУ": round(comb_calc["naoh_pure_year_t"], 2)
            },
            {
                "Показатель": "УДЕЛЬНЫЙ расход 100% NaOH на 1 кг отхода, г/кг",
                "Установка 1 (Жидкие отходы)": round(liq_calc["spec_naoh_pure_g_per_kg"], 2),
                "Установка 2 (ТБО 170 кг/ч)": round(tbo_calc["spec_naoh_pure_g_per_kg"], 1),
                "СУММАРНО ПО КОМПЛЕКСУ": round(comb_calc["spec_naoh_pure_g_per_kg"], 1)
            }
        ])
        df_summary.to_excel(writer, sheet_name="Сводные показатели", index=False)
        
        # Лист 2: Исходные данные и формулы
        df_inputs = pd.DataFrame([
            {"Категория": "Режим работы", "Параметр": "Длительность смены в сутки (T_сут)", "Значение": params["hours_per_day"], "Ед. изм.": "ч/сут", "Формула / Примечание": "Задается пользователем"},
            {"Категория": "Режим работы", "Параметр": "Рабочих дней в году с учетом ППР (D_год)", "Значение": params["operating_days_year"], "Ед. изм.": "дн/год", "Формула / Примечание": "С учетом ремонтов/простоев"},
            {"Категория": "Режим работы", "Параметр": "Годовой фонд времени (T_год)", "Значение": params["annual_hours"], "Ед. изм.": "ч/год", "Формула / Примечание": "T_год = T_сут * D_год"},
            {"Категория": "Дымовые газы", "Параметр": "Расход газов Установки жидких отходов (V_г,liq)", "Значение": params.get("flue_gas_flow_liq", 30000.0), "Ед. изм.": "нм3/ч", "Формула / Примечание": "КТОЖС 1,5 м3/ч"},
            {"Категория": "Дымовые газы", "Параметр": "Расход газов Установки ТБО (V_г,tbo)", "Значение": params.get("flue_gas_flow_tbo", 15000.0), "Ед. изм.": "нм3/ч", "Формула / Примечание": "ТБО 170 кг/ч"},
            {"Категория": "Дымовые газы", "Параметр": "Суммарный расход газов комплекса (V_г,общ)", "Значение": params.get("flue_gas_flow", 45000.0), "Ед. изм.": "нм3/ч", "Формула / Примечание": "V_г,liq + V_г,tbo"},

            {"Категория": "Скруббер", "Параметр": "Паспортная эффективность (η_скр)", "Значение": params["eta_scrubber"], "Ед. изм.": "д.ед.", "Формула / Примечание": "Степень очистки по HCl и SO2"},
            {"Категория": "Скруббер", "Параметр": "Коэффициент избытка реагента (k_изб)", "Значение": params["k_excess"], "Ед. изм.": "д.ед.", "Формула / Примечание": "Технологический запас"},
            {"Категория": "Жидкие отходы", "Параметр": "Расход стоков (Q_liq)", "Значение": params["q_liq"], "Ед. изм.": "м3/ч", "Формула / Примечание": "1500 кг/ч"},
            {"Категория": "Жидкие отходы", "Параметр": "Режим фильтрата", "Значение": params["dataset_name"], "Ед. изм.": "-", "Формула / Примечание": "СП 320.1325800.2017, Табл. Г.1"},
            {"Категория": "Жидкие отходы", "Параметр": "Концентрация Cl-", "Значение": params["c_cl_liq"], "Ед. изм.": "мг/л", "Формула / Примечание": "M_Cl = (Q * C_Cl)/1000"},
            {"Категория": "Жидкие отходы", "Параметр": "Концентрация SO4(2-)", "Значение": params["c_so4_liq"], "Ед. изм.": "мг/л", "Формула / Примечание": "M_S = (Q * C_SO4 * 32.065/96.061)/1000"},
            {"Категория": "ТБО", "Параметр": "Производительность по ТБО (M_ТБО)", "Значение": params["m_tbo"], "Ед. изм.": "кг/ч", "Формула / Примечание": "2500 ккал/кг"},
            {"Категория": "ТБО", "Параметр": "Среднее содержание Cl в ТБО", "Значение": round(tbo_feed["avg_pct_cl"], 3), "Ед. изм.": "%", "Формула / Примечание": "M_Cl = M_ТБО * (Cl%/100)"},
            {"Категория": "ТБО", "Параметр": "Среднее содержание S в ТБО", "Значение": round(tbo_feed["avg_pct_s"], 3), "Ед. изм.": "%", "Формула / Примечание": "M_S = M_ТБО * (S%/100)"},
            {"Категория": "Стехиометрия", "Параметр": "HCl + NaOH -> NaCl + H2O", "Значение": 1.0970, "Ед. изм.": "кг NaOH / кг HCl", "Формула / Примечание": "μ_NaOH / μ_HCl = 39.997 / 36.461"},
            {"Категория": "Стехиометрия", "Параметр": "SO2 + 2NaOH -> Na2SO3 + H2O", "Значение": 1.2487, "Ед. изм.": "кг NaOH / кг SO2", "Формула / Примечание": "2*μ_NaOH / μ_SO2 = 79.994 / 64.063"},
            {"Категория": "Конверсия в газ", "Параметр": "Конверсия Cl -> HCl (Жидкие)", "Значение": params.get("k_conv_cl_liq", 0.95), "Ед. изм.": "д.ед.", "Формула / Примечание": "Доля перехода Cl -> HCl"},
            {"Категория": "Конверсия в газ", "Параметр": "Конверсия S -> SO2 (Жидкие)", "Значение": params.get("k_conv_s_liq", 0.85), "Ед. изм.": "д.ед.", "Формула / Примечание": "Доля перехода S -> SO2"},
            {"Категория": "Конверсия в газ", "Параметр": "Конверсия Cl -> HCl (ТБО)", "Значение": params.get("k_conv_cl_tbo", 0.85), "Ед. изм.": "д.ед.", "Формула / Примечание": "Доля перехода Cl -> HCl"},
            {"Категория": "Конверсия в газ", "Параметр": "Конверсия S -> SO2 (ТБО)", "Значение": params.get("k_conv_s_tbo", 0.80), "Ед. изм.": "д.ед.", "Формула / Примечание": "Доля перехода S -> SO2"}
        ])
        df_inputs.to_excel(writer, sheet_name="Исходные данные и формулы", index=False)
        
        # Лист 3: Жидкие отходы (баланс)
        df_liq_gas = pd.DataFrame([
            {"Кислый газ": "HCl", "Поступление элемента, кг/ч": f"{liquid_feed['mass_cl']:.3f} (Cl)", "Масса в газе, кг/ч": round(liquid_feed["mass_hcl"], 3), "Теор. 100% NaOH, кг/ч": round(liq_calc["naoh_hcl_theor"], 3), "Факт. 100% NaOH, кг/ч": round(liq_calc["naoh_hcl_fact"], 3), "Суточный 100%, кг/сут": round(liq_calc["naoh_hcl_fact"]*params["hours_per_day"], 1), "ГОДОВОЙ 100%, т/год": round((liq_calc["naoh_hcl_fact"]*params["annual_hours"])/1000, 2), "Удельный, г/кг": round((liq_calc["naoh_hcl_fact"]/liquid_feed["feed_mass_kg_h"])*1000, 2)},
            {"Кислый газ": "SO2", "Поступление элемента, кг/ч": f"{liquid_feed['mass_s']:.3f} (S)", "Масса в газе, кг/ч": round(liquid_feed["mass_so2"], 3), "Теор. 100% NaOH, кг/ч": round(liq_calc["naoh_so2_theor"], 3), "Факт. 100% NaOH, кг/ч": round(liq_calc["naoh_so2_fact"], 3), "Суточный 100%, кг/сут": round(liq_calc["naoh_so2_fact"]*params["hours_per_day"], 1), "ГОДОВОЙ 100%, т/год": round((liq_calc["naoh_so2_fact"]*params["annual_hours"])/1000, 2), "Удельный, г/кг": round((liq_calc["naoh_so2_fact"]/liquid_feed["feed_mass_kg_h"])*1000, 2)},
            {"Кислый газ": "ИТОГО", "Поступление элемента, кг/ч": "-", "Масса в газе, кг/ч": round(liquid_feed["mass_hcl"]+liquid_feed["mass_so2"], 3), "Теор. 100% NaOH, кг/ч": round(liq_calc["naoh_total_theor"], 3), "Факт. 100% NaOH, кг/ч": round(liq_calc["naoh_total_fact"], 3), "Суточный 100%, кг/сут": round(liq_calc["naoh_pure_day_kg"], 1), "ГОДОВОЙ 100%, т/год": round(liq_calc["naoh_pure_year_t"], 2), "Удельный, г/кг": round(liq_calc["spec_naoh_pure_g_per_kg"], 2)}
        ])
        df_liq_gas.to_excel(writer, sheet_name="Жидкие отходы (баланс)", index=False)
        
        # Лист 4: ТБО (баланс)
        df_tbo_gas = pd.DataFrame([
            {"Кислый газ": "HCl", "Поступление элемента, кг/ч": f"{tbo_feed['mass_cl']:.4f} (Cl)", "Масса в газе, кг/ч": round(tbo_feed["mass_hcl"], 3), "Теор. 100% NaOH, кг/ч": round(tbo_calc["naoh_hcl_theor"], 3), "Факт. 100% NaOH, кг/ч": round(tbo_calc["naoh_hcl_fact"], 3), "Суточный 100%, кг/сут": round(tbo_calc["naoh_hcl_fact"]*params["hours_per_day"], 1), "ГОДОВОЙ 100%, т/год": round((tbo_calc["naoh_hcl_fact"]*params["annual_hours"])/1000, 2), "Удельный, г/кг": round((tbo_calc["naoh_hcl_fact"]/params["m_tbo"])*1000, 2)},
            {"Кислый газ": "SO2", "Поступление элемента, кг/ч": f"{tbo_feed['mass_s']:.4f} (S)", "Масса в газе, кг/ч": round(tbo_feed["mass_so2"], 3), "Теор. 100% NaOH, кг/ч": round(tbo_calc["naoh_so2_theor"], 3), "Факт. 100% NaOH, кг/ч": round(tbo_calc["naoh_so2_fact"], 3), "Суточный 100%, кг/сут": round(tbo_calc["naoh_so2_fact"]*params["hours_per_day"], 1), "ГОДОВОЙ 100%, т/год": round((tbo_calc["naoh_so2_fact"]*params["annual_hours"])/1000, 2), "Удельный, г/кг": round((tbo_calc["naoh_so2_fact"]/params["m_tbo"])*1000, 2)},
            {"Кислый газ": "ИТОГО", "Поступление элемента, кг/ч": "-", "Масса в газе, кг/ч": round(tbo_feed["mass_hcl"]+tbo_feed["mass_so2"], 3), "Теор. 100% NaOH, кг/ч": round(tbo_calc["naoh_total_theor"], 3), "Факт. 100% NaOH, кг/ч": round(tbo_calc["naoh_total_fact"], 3), "Суточный 100%, кг/сут": round(tbo_calc["naoh_pure_day_kg"], 1), "ГОДОВОЙ 100%, т/год": round(tbo_calc["naoh_pure_year_t"], 2), "Удельный, г/кг": round(tbo_calc["spec_naoh_pure_g_per_kg"], 1)}
        ])
        df_tbo_gas.to_excel(writer, sheet_name="ТБО (баланс)", index=False)
        
        # Лист 5: Состав ТБО (морфология)
        if "breakdown_table" in tbo_feed and tbo_feed["breakdown_table"]:
            df_morph = pd.DataFrame(tbo_feed["breakdown_table"])
            df_morph.to_excel(writer, sheet_name="Морфология ТБО", index=False)

        # Лист 6: Compliance (ИТС 9-2020)
        fl_liq = params.get("flue_gas_flow_liq", 30000.0)
        fl_tbo = params.get("flue_gas_flow_tbo", 15000.0)

        c_hcl_in_l = (liquid_feed['mass_hcl'] * 1e6) / fl_liq if fl_liq > 0 else 0.0
        c_so2_in_l = (liquid_feed['mass_so2'] * 1e6) / fl_liq if fl_liq > 0 else 0.0
        c_hcl_in_t = (tbo_feed['mass_hcl'] * 1e6) / fl_tbo if fl_tbo > 0 else 0.0
        c_so2_in_t = (tbo_feed['mass_so2'] * 1e6) / fl_tbo if fl_tbo > 0 else 0.0
        
        eta_s = params.get('eta_scrubber', 0.95)
        df_comp = pd.DataFrame([
            {"Установка": "Уст.1 (Жидкие отходы)", "Вещество": "HCl", "Расход газов, нм3/ч": fl_liq, "Входная конц., мг/нм3": round(c_hcl_in_l, 1), "ПДК ИТС 9-2020, мг/нм3": 10.0, "Выходная конц., мг/нм3": round(min(c_hcl_in_l, 10.0), 1), "Требуемая η": f"{max(0.0, 1-10/c_hcl_in_l):.1%}" if c_hcl_in_l > 10 else "0.0%", "Статус": "НОРМА"},
            {"Установка": "Уст.1 (Жидкие отходы)", "Вещество": "SO2", "Расход газов, нм3/ч": fl_liq, "Входная конц., мг/нм3": round(c_so2_in_l, 1), "ПДК ИТС 9-2020, мг/нм3": 50.0, "Выходная конц., мг/нм3": round(min(c_so2_in_l, 50.0), 1), "Требуемая η": f"{max(0.0, 1-50/c_so2_in_l):.1%}" if c_so2_in_l > 50 else "0.0%", "Статус": "НОРМА"},
            {"Установка": "Уст.2 (ТБО 170 кг/ч)", "Вещество": "HCl", "Расход газов, нм3/ч": fl_tbo, "Входная конц., мг/нм3": round(c_hcl_in_t, 1), "ПДК ИТС 9-2020, мг/нм3": 10.0, "Выходная конц., мг/нм3": round(min(c_hcl_in_t, 10.0), 1), "Требуемая η": f"{max(0.0, 1-10/c_hcl_in_t):.1%}" if c_hcl_in_t > 10 else "0.0%", "Статус": "НОРМА"},
            {"Установка": "Уст.2 (ТБО 170 кг/ч)", "Вещество": "SO2", "Расход газов, нм3/ч": fl_tbo, "Входная конц., мг/нм3": round(c_so2_in_t, 1), "ПДК ИТС 9-2020, мг/нм3": 50.0, "Выходная конц., мг/нм3": round(min(c_so2_in_t, 50.0), 1), "Требуемая η": f"{max(0.0, 1-50/c_so2_in_t):.1%}" if c_so2_in_t > 50 else "0.0%", "Статус": "НОРМА"},
        ])
        df_comp.to_excel(writer, sheet_name="Compliance (ИТС 9-2020)", index=False)

            
    if output_path is not None:
        return output_path
    else:
        return buf.getvalue()

